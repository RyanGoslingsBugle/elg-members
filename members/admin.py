from datetime import datetime, timedelta
from django.conf.urls import url
from django.contrib import admin
from django.http import HttpResponse
from django.core.exceptions import PermissionDenied
from django.utils.encoding import smart_str
import csv
from django.utils.six import BytesIO
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm

from .models import Member

admin.site.site_header = 'ELG Member Admin'
admin.site.index_title = 'Member Admin'

@admin.register(Member)
class MemberAdmin(admin.ModelAdmin):
    list_display = ('last_name', 'first_name', 'current_membership', 'card_number', 'email')
    list_filter = ['current_membership']
    fieldsets = (
        (None, {'fields': ('last_name', 'first_name', 'email', 'address', 'phone_number')}),
        ('Membership', {'fields': ('current_membership', 'card_number')}),
    )

    def get_urls(self):
        urls = super(MemberAdmin, self).get_urls()
        custom_urls = [
            url(r'^export', self.export, name="export"),
            url(r'^get_door', self.get_door, name="get_door")
        ]
        return urls + custom_urls

    @staticmethod
    def export(request):
        return download_csv(request, Member.objects.all())

    @staticmethod
    def get_door(request):
        return make_door(request, Member.objects.all().order_by('last_name'))


def download_csv(request, queryset):
    if not request.user.is_staff:
        raise PermissionDenied

    opts = queryset.model._meta
    response = HttpResponse(content_type='text/csv; name=%s' % smart_str("export.csv"))
    # force download
    response['Content-Disposition'] = 'attachment; filename=%s' % smart_str("export.csv")
    # the csv writer
    writer = csv.writer(response)
    field_names = [field.name for field in opts.fields]
    # Write a first row with header information
    writer.writerow(field_names)
    # Write data rows
    for obj in queryset:
        writer.writerow([getattr(obj, field) for field in field_names])

    return response


def make_door(request, queryset):

    if not request.user.is_staff:
        raise PermissionDenied

    if request.GET['start_date'] is None:
        raise KeyError('Request must include start_date')

    start_date = datetime.strptime(request.GET['start_date'], "%m/%d/%Y")

    document = Document('template.docx')

    # write header values
    hd_table = document.sections[0].header.tables[0]
    cells = hd_table.rows[1].cells
    for count, i in enumerate(range(2, 8)):
        days = timedelta(days=7*count)
        text = (start_date + days).strftime("%d %b")
        par = cells[i].paragraphs[0]
        par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        par.text = text

    # add table
    table = document.tables[0]
    table.style = document.styles["Table Grid"]

    # write rows for members
    for count, obj in enumerate(queryset):
        if len(table.rows) > count:
            row = table.rows[count]
        else:
            row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Cm(0.7)
        cells = row.cells
        text = obj.first_name + " " + obj.last_name
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        par = cells[0].paragraphs[0]
        par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        par.add_run(text)
        if obj.card_number < 999:
            if obj.current_membership:
                text = str(obj.card_number)
            else:
                text = "[%s]" % obj.card_number
            par = cells[1].paragraphs[0]
            par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            par.add_run(text)

    for i in range(30):
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = Cm(0.7)

    target_stream = BytesIO()
    document.save(target_stream)
    target_stream.seek(0)

    response = HttpResponse(target_stream,
                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document; '
                                         'name=%s' % smart_str("door.docx"))
    response['Content-Disposition'] = 'attachment; filename=%s' % smart_str("door.docx")

    return response
